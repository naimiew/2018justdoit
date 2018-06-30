# -*- coding: utf-8 -*-
#场景：目前职位测试，需按文档校验页面字段字符长度（边界值)
#从excel文档中获取字符长度限制，生成方便检查的1234567890+的整数
#保存于word中，方便测试结果  （excel中大于15位整数会使文件损坏）

# 读写高版本 excel
import openpyxl
#创建并写入word文档
import docx
import datetime
#创建内存中的word文档对象
#make_docx = docx.Document('D:/123/实战/excel取数存储到word/save.docx')
make_docx = docx.Document()

def get_num(Data2int):
    # 定义一个空格字符串
    s = ''
    # excel取值做最大值
    num = Data2int
    # 定义初始整数
    n = 0
    # 定义阈值
    m = 1
    # 循环制造整数
    while m <= num:
        m = m + 1
        n = n + 1
        n = n % 10
        t = str(n)
        s = s + t
        # print(s)
    #print(s)
    return s

#excel路径  ‪C:/123/dk.xlsx
excel_path = 'c:/123/'
excel_name = 'dk2.xlsx'

#打开文件：
rw_excel = openpyxl.load_workbook(excel_path+excel_name)
#查找文件判断
if rw_excel != '':
    print('已找到excel文件，文件名为@@@@'+excel_name)
else:
    print('查找excel失败''！！！！！！！！')
#获取sheet：
work_sheet = rw_excel.worksheets[0] #通过表名获取
print('该sheet页为'+work_sheet.title)

  #  read_excel.get_sheet_by_name("Sheet2")  建议不适用这个过时的方法
#获取行数和列数：
rows = work_sheet.max_row  #获取行数
cols = work_sheet.max_column    #获取列数
print('rows_num:'+str(rows)+'\n'+'cols_num:'+str(cols))
#设置启动行
go_row = 2
#设置动态关联excel行列
while go_row <= rows:
    #执行excel行
    go_row = go_row + 1
    #获取单元格值：
    Data1 = work_sheet.cell(go_row, column=1).value
    Data2 = work_sheet.cell(go_row, column=2).value
    Data3 = work_sheet.cell(go_row, column=3).value
    Data4 = work_sheet.cell(go_row, column=4).value
    text_Data = str(Data1)+'##'+str(Data2)+'##'+str(Data3)+'##'+str(Data4)
    print(text_Data)
    Data = work_sheet.cell(go_row,column=8).value  #获取表格内容，是从第一行第一列是从1开始的，注意不要丢掉 .value
    #print(Data)
    #获取数据判断 if not 判空与真
    if Data:
        print('获取数据成功！！！！！！！！')
    else:
        print('获取数据失败！！！！！！！！')

    #print(Data) 有机会可以改if判断
    if isinstance(Data,str):
        Data2int = int(Data) #str转int
    else:
        Data2int = Data #str直传

    #print(Data2int)        # 有机会可以改if判断
    #判断结束
    if not Data2int:
        print('获取数据结束！！！！！！！！！！！！！！！')
        break

    #第一次循环Data2int
    s1 = get_num(Data2int)
    #print(s1)
    #写入数据 def cell(self, row, column, value=None): 新版本
    #work_sheet.cell(go_row,column=9,value=str(s1))
    #换word存储
    make_docx.add_paragraph(text=text_Data)
    make_docx.add_paragraph(text='第'+str(go_row)+'行')
    make_docx.add_paragraph(text=str(s1))
    print('写入成功第'+str(go_row)+'行数据:'+str(s1))
    #取数据校验  缺判断
    #new_Data1 = work_sheet.cell(go_row,column=9).value
    #print('写入数据为@@@@@@'+str(go_row)+'列;数据为'+str(new_Data1))
    # 保存excel
    #rw_excel.save(excel_path + excel_name)
    #print('写入数据为@@@@@@第'+str(go_row)+'保存成功')
    #换word保存文件
#    t = datetime.datetime.now().strftime('%Y%m%d%H%M%S')
#    save_path = 'D:/123/实战/excel取数存储到word/save'+t+'.docx'
#    make_docx.save(save_path)

    #第二次循环Data2int+1
    s2 = get_num(Data2int+1)
    #print(s2)
    #写入数据
    #work_sheet.cell(go_row,column=10,value=str(s2))
    #换word存储
    make_docx.add_paragraph(text=str(s2))
    print('写入成功第' + str(go_row) + '行数据:' + str(s2))
    #取数据校验  缺判断
    #new_Data2 = work_sheet.cell(go_row,column=10).value
    #print('写入数据为@@@@@@第'+str(go_row)+'列;数据为'+str(new_Data2))
    #保存excel
    #rw_excel.save(excel_path+excel_name)
    #print('写入数据为@@@@@@第'+str(go_row)+'保存成功')


#拼贴完整路径
excel_path_full = excel_path+excel_name
print(excel_path_full)
#保存excel
#rw_excel.save(excel_path+excel_name)
#print('保存成功'+excel_path_full)
# 换word保存文件
t = datetime.datetime.now().strftime('%Y%m%d%H%M%S')
save_path = 'D:/123/实战/excel取数存储到word/save' + t + '.docx'
make_docx.save(save_path)
print(save_path+'保存成功!!!!')
